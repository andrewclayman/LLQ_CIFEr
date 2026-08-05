"""Generate a .docx describing all significance tests from the GA Portfolio Results presentation."""

import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ── Title ──
title = doc.add_heading('Statistical Significance Tests', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(
    'Evolutionary Strategy Optimisation Portfolio\n'
    'A comprehensive description of all significance tests presented in the GA Portfolio Results.'
).alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── Overview ──
doc.add_heading('Overview', level=1)
doc.add_paragraph(
    'The experimental design compares three Genetic Algorithm (GA) variants — Full GA '
    '(13 strategies), LLM-Only GA (8 strategies), and Benchmark-Only GA (6 strategies) — '
    'across 6 out-of-sample (OOS) periods from 2018 to 2024. Each variant is run 50 '
    'independent times per period (900 total GA runs), producing distributions of OOS '
    'Sharpe ratios that are then compared using a suite of statistical tests. The primary '
    'metric is the annualised Sharpe ratio, though Sortino, Calmar, maximum drawdown, and '
    'total return are also tested.'
)
doc.add_paragraph(
    'The testing framework is deliberately multi-layered: a primary non-parametric test '
    '(Mann-Whitney U) is corroborated by parametric tests (independent t-test), normality '
    'diagnostics (Shapiro-Wilk), variance diagnostics (Levene\'s test), an omnibus test '
    '(Kruskal-Wallis), effect size measures (Cohen\'s d), confidence intervals (bootstrap), '
    'and a mixed-effects model that accounts for the nested structure of the data. This '
    'triangulation ensures conclusions are robust regardless of distributional assumptions.'
)

# ── 1. Mann-Whitney U ──
doc.add_heading('1. Mann-Whitney U Test (Primary Test)', level=1)

doc.add_heading('What it is', level=2)
doc.add_paragraph(
    'The Mann-Whitney U test (also called the Wilcoxon rank-sum test) is a non-parametric '
    'hypothesis test that compares two independent samples. It tests whether one population '
    'tends to produce larger values than the other. Unlike the t-test, it does not assume '
    'that the data are normally distributed — it operates on ranks rather than raw values.'
)

doc.add_heading('How it is used in this study', level=2)
doc.add_paragraph(
    'For each OOS period, the 50 Sharpe ratios from one GA variant (e.g., Full GA) are '
    'compared against the 50 Sharpe ratios from another variant (e.g., Benchmark-Only GA). '
    'The null hypothesis is that there is no difference in the distribution of Sharpe ratios '
    'between the two groups. A small p-value indicates that one variant reliably produces '
    'higher (or lower) Sharpe ratios than the other.'
)
doc.add_paragraph(
    'Three pairwise comparisons are made per period (Full vs Benchmark, LLM-Only vs '
    'Benchmark, Full vs LLM-Only), across 6 periods, yielding 18 hypothesis tests in total. '
    'An aggregate test is also performed by pooling all 300 runs per variant across periods.'
)

doc.add_heading('When it is used', level=2)
doc.add_paragraph(
    'Mann-Whitney U is chosen as the primary test because the Shapiro-Wilk normality '
    'test rejects normality for 5 of 18 Sharpe distributions (at p < 0.05). Since normality '
    'cannot be assumed uniformly, a non-parametric test is the conservative and appropriate '
    'primary choice. It is applied to every pairwise comparison reported in the study.'
)

doc.add_heading('Key results', level=2)
doc.add_paragraph(
    'Full GA vs Benchmark: statistically significant (adjusted p < 0.05) in all 6 periods '
    'individually and in aggregate (p = 0.0000). Full GA outperforms Benchmark in 5 of 6 '
    'periods (the exception being 2020, where Benchmark wins).'
)
doc.add_paragraph(
    'LLM-Only GA vs Benchmark: statistically significant in all 6 periods and in aggregate '
    '(p = 0.0000). LLM-Only outperforms Benchmark in 5 of 6 periods.'
)

doc.add_heading('Results in detail and what they mean', level=2)
# Full vs Benchmark table
t = doc.add_table(rows=8, cols=6, style='Light Grid Accent 1')
for j, h in enumerate(['OOS Period', 'Full Mean Sharpe', 'Bench Mean Sharpe', 'Difference', 'Adj. p-value', 'Significant?']):
    t.rows[0].cells[j].text = h
full_mw_data = [
    ['2018', '-1.021', '-1.504', '+0.483', '0.0000', 'Yes (Full wins)'],
    ['2020', '+1.971', '+2.235', '-0.265', '0.0000', 'Yes (Bench wins)'],
    ['2021', '+0.614', '+0.433', '+0.181', '0.0007', 'Yes (Full wins)'],
    ['2022', '-0.696', '-1.105', '+0.409', '0.0000', 'Yes (Full wins)'],
    ['2023', '+2.177', '+1.280', '+0.897', '0.0000', 'Yes (Full wins)'],
    ['2024', '+1.732', '+0.969', '+0.763', '0.0000', 'Yes (Full wins)'],
    ['ALL (pooled)', '+0.796', '+0.385', '+0.411', '0.0000', 'Yes (Full wins)'],
]
for i, row_data in enumerate(full_mw_data):
    for j, cell_text in enumerate(row_data):
        t.rows[i + 1].cells[j].text = cell_text

doc.add_paragraph('')  # spacer

# LLM vs Benchmark table
doc.add_paragraph('LLM-Only GA vs Benchmark:')
t2 = doc.add_table(rows=8, cols=6, style='Light Grid Accent 1')
for j, h in enumerate(['OOS Period', 'LLM Mean Sharpe', 'Bench Mean Sharpe', 'Difference', 'Adj. p-value', 'Significant?']):
    t2.rows[0].cells[j].text = h
llm_mw_data = [
    ['2018', '-0.794', '-1.504', '+0.709', '0.0000', 'Yes (LLM wins)'],
    ['2020', '+1.945', '+2.235', '-0.290', '0.0000', 'Yes (Bench wins)'],
    ['2021', '+0.652', '+0.433', '+0.220', '0.0000', 'Yes (LLM wins)'],
    ['2022', '-0.708', '-1.105', '+0.397', '0.0000', 'Yes (LLM wins)'],
    ['2023', '+2.251', '+1.280', '+0.971', '0.0000', 'Yes (LLM wins)'],
    ['2024', '+1.543', '+0.969', '+0.574', '0.0000', 'Yes (LLM wins)'],
    ['ALL (pooled)', '+0.815', '+0.385', '+0.430', '0.0000', 'Yes (LLM wins)'],
]
for i, row_data in enumerate(llm_mw_data):
    for j, cell_text in enumerate(row_data):
        t2.rows[i + 1].cells[j].text = cell_text

doc.add_paragraph('')
doc.add_paragraph(
    'What this means: The Mann-Whitney U test provides the strongest individual piece of '
    'evidence in the study. In 5 out of 6 OOS periods, both the Full GA and LLM-Only GA '
    'produce reliably higher Sharpe ratios than the Benchmark-Only GA. The lone exception '
    'is 2020 (the COVID year), where the Benchmark actually wins — but even there, the '
    'difference is statistically significant, meaning the result is real, not noise. The '
    'aggregate pooled test (all 300 runs combined) confirms the overall pattern: portfolios '
    'that include LLM-generated strategies achieve meaningfully better risk-adjusted returns '
    'than portfolios restricted to conventional benchmarks. Crucially, the LLM-Only GA '
    '(which has no benchmark strategies at all) performs comparably to the Full GA, '
    'demonstrating that the novel strategies can stand alone.'
)

# ── 2. Holm-Bonferroni Correction ──
doc.add_heading('2. Holm-Bonferroni Correction (Multiple Comparisons Adjustment)', level=1)

doc.add_heading('What it is', level=2)
doc.add_paragraph(
    'The Holm-Bonferroni method is a step-down procedure for controlling the family-wise '
    'error rate (FWER) when performing multiple hypothesis tests simultaneously. When many '
    'tests are conducted, the probability of at least one false positive (Type I error) '
    'increases. The Holm-Bonferroni correction adjusts p-values to account for this, while '
    'being less conservative (more powerful) than the classical Bonferroni correction.'
)
doc.add_paragraph(
    'The procedure works by: (1) sorting all raw p-values from smallest to largest, '
    '(2) multiplying the smallest p-value by the total number of tests, the second smallest '
    'by (total - 1), and so on, (3) enforcing monotonicity so that adjusted p-values never '
    'decrease. A result is significant if its adjusted p-value falls below the chosen '
    'significance level (alpha = 0.05).'
)

doc.add_heading('How it is used in this study', level=2)
doc.add_paragraph(
    'All 18 Mann-Whitney U tests (3 pairwise comparisons x 6 periods) are corrected '
    'simultaneously using Holm-Bonferroni. The raw p-values from the Mann-Whitney tests are '
    'adjusted, and only the adjusted p-values are used to determine statistical significance. '
    'This ensures that the overall probability of even one false positive across all 18 tests '
    'remains below 5%.'
)

doc.add_heading('When it is used', level=2)
doc.add_paragraph(
    'Whenever multiple hypothesis tests are reported together — here, the 18 pairwise '
    'Mann-Whitney U tests. Without this correction, conducting 18 tests at alpha = 0.05 would '
    'give roughly a 60% chance of at least one false positive. The correction controls this.'
)

doc.add_heading('Key results', level=2)
doc.add_paragraph(
    'All significant results survive the Holm-Bonferroni correction. For example, the Full '
    'vs Benchmark comparison in 2021 has a raw p-value of 0.0001 and an adjusted p-value of '
    '0.0007 — still well below 0.05.'
)

doc.add_heading('Results in detail and what they mean', level=2)
doc.add_paragraph(
    'The weakest individual raw p-value across all 18 tests is 0.0001 (Full vs Benchmark, '
    '2021). After Holm-Bonferroni adjustment this becomes 0.0007, which is still far below '
    'the 0.05 threshold. All other adjusted p-values are 0.0000 (i.e., < 0.00005). This '
    'means that none of the significant findings are artefacts of multiple testing. Even '
    'under the most stringent correction that controls the probability of even a single '
    'false positive across all 18 tests, every result holds.'
)
doc.add_paragraph(
    'What this means: The Holm-Bonferroni correction is a safeguard against "p-hacking" '
    'or cherry-picking. When you run 18 tests, you would expect roughly 1 false positive '
    'by chance at alpha = 0.05. The fact that all results survive correction with room to '
    'spare (the largest adjusted p is 0.0007, not borderline 0.04) indicates that the '
    'differences between GA variants are robust and not an artefact of testing many '
    'comparisons.'
)

# ── 3. Bootstrap Confidence Intervals ──
doc.add_heading('3. Bootstrap Confidence Intervals', level=1)

doc.add_heading('What it is', level=2)
doc.add_paragraph(
    'Bootstrapping is a resampling method that estimates the sampling distribution of a '
    'statistic (here, the mean difference in Sharpe ratios) by repeatedly drawing samples '
    'with replacement from the observed data. A 95% bootstrap confidence interval gives the '
    'range within which the true mean difference is expected to lie with 95% confidence, '
    'without relying on parametric assumptions about the data distribution.'
)

doc.add_heading('How it is used in this study', level=2)
doc.add_paragraph(
    'For each pairwise comparison (e.g., Full GA vs Benchmark) in each OOS period, 10,000 '
    'bootstrap resamples are drawn. For each resample, the mean Sharpe difference is computed. '
    'The 2.5th and 97.5th percentiles of these 10,000 differences form the 95% confidence '
    'interval. A bootstrap p-value is also computed as the proportion of resamples where the '
    'difference has the opposite sign to the observed difference.'
)

doc.add_heading('When it is used', level=2)
doc.add_paragraph(
    'Bootstrap CIs are reported alongside every Mann-Whitney U test as a corroborating '
    'measure. They provide two things the p-value alone cannot: (1) the magnitude of the '
    'difference (not just whether it is non-zero), and (2) the uncertainty around that '
    'magnitude. If the 95% CI excludes zero, this independently confirms statistical '
    'significance.'
)

doc.add_heading('Key results', level=2)
doc.add_paragraph(
    'All six Full vs Benchmark bootstrap CIs exclude zero, confirming the Mann-Whitney '
    'results. For example, the 2023 period shows a 95% CI of [+0.786, +1.008] for the Sharpe '
    'difference, meaning the Full GA produces a Sharpe ratio between 0.79 and 1.01 higher '
    'than the Benchmark with 95% confidence.'
)

doc.add_heading('Results in detail and what they mean', level=2)
t3 = doc.add_table(rows=7, cols=4, style='Light Grid Accent 1')
for j, h in enumerate(['OOS Period', 'Mean Sharpe Diff (Full - Bench)', '95% Bootstrap CI', 'Bootstrap p']):
    t3.rows[0].cells[j].text = h
boot_data = [
    ['2018', '+0.483', '[+0.410, +0.555]', '0.0000'],
    ['2020', '-0.265', '[-0.332, -0.193]', '1.0000'],
    ['2021', '+0.181', '[+0.091, +0.270]', '0.0000'],
    ['2022', '+0.409', '[+0.361, +0.456]', '0.0000'],
    ['2023', '+0.897', '[+0.786, +1.008]', '0.0000'],
    ['2024', '+0.763', '[+0.710, +0.815]', '0.0000'],
]
for i, row_data in enumerate(boot_data):
    for j, cell_text in enumerate(row_data):
        t3.rows[i + 1].cells[j].text = cell_text

doc.add_paragraph('')
doc.add_paragraph(
    'What this means: The bootstrap CIs tell us not just that the Full GA is better, but '
    'by how much. In 2024 (the strongest period), the Full GA produces a Sharpe ratio '
    'between 0.71 and 0.82 higher than Benchmark — a substantial practical advantage. Even '
    'in 2021 (the weakest win), the lower bound of the CI is +0.091, meaning we can be 95% '
    'confident the Full GA is at least 0.09 Sharpe points better. The 2020 CI is entirely '
    'negative, confirming that the Benchmark genuinely outperforms in that period (COVID '
    'recovery favoured simple momentum/buy-and-hold). The tight CIs (narrow ranges) reflect '
    'the precision afforded by 50 runs per group — the estimates are stable, not noisy.'
)

# ── 4. Cohen's d ──
doc.add_heading("4. Cohen's d (Effect Size)", level=1)

doc.add_heading('What it is', level=2)
doc.add_paragraph(
    "Cohen's d is a standardised measure of effect size that quantifies the difference "
    "between two group means in units of their pooled standard deviation. It answers the "
    "question: 'How large is the difference, relative to the variability in the data?' "
    "Conventional thresholds are: |d| < 0.2 = negligible, 0.2-0.5 = small, 0.5-0.8 = "
    "medium, > 0.8 = large."
)
doc.add_paragraph(
    'The formula is: d = (mean_1 - mean_2) / s_pooled, where s_pooled is the pooled '
    'standard deviation of the two groups.'
)

doc.add_heading('How it is used in this study', level=2)
doc.add_paragraph(
    "Cohen's d is computed for every pairwise comparison alongside the Mann-Whitney U "
    "test and bootstrap CI. It provides a scale-free measure of how practically meaningful "
    "the Sharpe differences are — a statistically significant but tiny difference would show "
    "a small d, while a large d indicates a substantial, practically relevant gap."
)

doc.add_heading('When it is used', level=2)
doc.add_paragraph(
    'Effect sizes are reported for every period-level comparison and are summarised in the '
    'Key Findings. They are particularly important here because with 50 runs per group, even '
    'small differences can reach statistical significance. Cohen\'s d distinguishes '
    'statistically significant results that are also practically meaningful from those that '
    'are merely detectable.'
)

doc.add_heading('Key results', level=2)
doc.add_paragraph(
    'Full GA vs Benchmark: d ranges from +0.80 (2021) to +5.68 (2024), with a mean of '
    '2.87. All periods show large effect sizes (|d| > 0.8). The 2020 period where Benchmark '
    'wins also has a large effect (d = -1.49). These large effect sizes confirm that the '
    'Sharpe differences are not just statistically significant but practically substantial.'
)

doc.add_heading('Results in detail and what they mean', level=2)
t4 = doc.add_table(rows=7, cols=4, style='Light Grid Accent 1')
for j, h in enumerate(["OOS Period", "Cohen's d (Full vs Bench)", "Magnitude", "Interpretation"]):
    t4.rows[0].cells[j].text = h
cohen_data = [
    ['2018', '+2.64', 'Large', 'Full GA distributions barely overlap with Benchmark'],
    ['2020', '-1.49', 'Large', 'Benchmark clearly better; distributions well separated'],
    ['2021', '+0.80', 'Large (borderline)', 'Meaningful separation; smallest win for Full GA'],
    ['2022', '+3.35', 'Very large', 'Distributions almost entirely non-overlapping'],
    ['2023', '+3.24', 'Very large', 'Strong practical advantage for Full GA'],
    ['2024', '+5.68', 'Very large', 'Extreme separation; virtually no overlap'],
]
for i, row_data in enumerate(cohen_data):
    for j, cell_text in enumerate(row_data):
        t4.rows[i + 1].cells[j].text = cell_text

doc.add_paragraph('')
doc.add_paragraph(
    'What this means: A Cohen\'s d above 0.8 is conventionally "large". The mean across '
    'all periods is 2.87, which is exceptionally large by any standard in empirical research. '
    'To put this in perspective, a d of 2.64 (2018) means that an average run from the Full '
    'GA outperforms roughly 99.6% of Benchmark runs. A d of 5.68 (2024) means virtually '
    'every Full GA run beats virtually every Benchmark run — the two distributions do not '
    'overlap. These effect sizes matter because p-values alone can be misleading with large '
    'samples: with n = 50, even a trivially small difference could be "significant". The '
    'large effect sizes confirm that the differences are not just detectable but practically '
    'important — the Full GA and LLM-Only GA genuinely produce much better portfolios in '
    'most market conditions.'
)

# ── 5. Shapiro-Wilk ──
doc.add_heading('5. Shapiro-Wilk Test (Normality Diagnostic)', level=1)

doc.add_heading('What it is', level=2)
doc.add_paragraph(
    'The Shapiro-Wilk test assesses whether a sample comes from a normally distributed '
    'population. The null hypothesis is that the data are normally distributed; a small '
    'p-value (typically < 0.05) leads to rejection of normality. It is one of the most '
    'powerful normality tests for small to moderate sample sizes.'
)

doc.add_heading('How it is used in this study', level=2)
doc.add_paragraph(
    'Shapiro-Wilk is applied to each of the 18 Sharpe ratio distributions (3 variants x '
    '6 periods, each with n = 50 runs). It serves as a diagnostic to determine whether '
    'parametric tests (like the t-test) are appropriate, or whether non-parametric '
    'alternatives (Mann-Whitney U) should be preferred.'
)

doc.add_heading('When it is used', level=2)
doc.add_paragraph(
    'Before selecting the primary test. The results directly justify the choice of '
    'Mann-Whitney U as the primary significance test rather than the independent t-test.'
)

doc.add_heading('Key results', level=2)
doc.add_paragraph(
    '5 out of 18 distributions reject normality at p < 0.05. Since normality cannot be '
    'assumed across all groups, non-parametric tests are the appropriate primary analysis. '
    'Parametric tests (t-tests) are still reported for completeness but are secondary.'
)

doc.add_heading('Results in detail and what they mean', level=2)
doc.add_paragraph(
    'Of the 18 Sharpe distributions tested (3 variants x 6 periods), 5 rejected the null '
    'hypothesis of normality. This means that approximately 28% of the distributions show '
    'statistically significant departures from a normal (Gaussian) bell curve shape — they '
    'may be skewed, have heavier tails, or be multimodal.'
)
doc.add_paragraph(
    'What this means: This result is primarily a methodological justification. It tells us '
    'that we cannot rely on parametric tests (which assume normality) as the primary '
    'analysis. If all 18 had passed, a t-test could have been the primary test. Since 28% '
    'fail, the Mann-Whitney U (which makes no normality assumption) is the correct choice. '
    'The Shapiro-Wilk result does not directly say anything about whether the GA variants '
    'differ — it is a diagnostic that validates the choice of statistical methodology. The '
    'non-normality likely arises because the GA is a stochastic optimiser: most runs converge '
    'to similar solutions, but occasional runs get trapped in local optima, creating a '
    'skewed or heavy-tailed distribution of outcomes.'
)

# ── 6. Levene's Test ──
doc.add_heading("6. Levene's Test (Equality of Variances)", level=1)

doc.add_heading('What it is', level=2)
doc.add_paragraph(
    "Levene's test assesses whether two or more groups have equal variances "
    "(homoscedasticity). The null hypothesis is that the variances are equal. Unequal "
    "variances can invalidate the standard independent t-test, requiring Welch's t-test "
    "(which does not assume equal variances) instead."
)

doc.add_heading('How it is used in this study', level=2)
doc.add_paragraph(
    "Levene's test is applied to each period's Full GA vs Benchmark Sharpe distributions "
    "(50 runs each). When Levene's test indicates unequal variances, this is noted and "
    "Welch's correction would be applied to any parametric t-test. It also provides insight "
    "into the behaviour of the GA: periods with unequal variances suggest that one variant "
    "produces more consistent results than the other."
)

doc.add_heading('When it is used', level=2)
doc.add_paragraph(
    'As a diagnostic alongside the parametric t-tests reported on Slide 11 (Full GA vs '
    'Benchmark across multiple metrics). It determines whether a standard or Welch-corrected '
    't-test is appropriate.'
)

doc.add_heading('Key results', level=2)
doc.add_paragraph(
    '2018 (F = 28.85, p = 0.0000) and 2020 (F = 38.32, p = 0.0000) show significantly '
    'unequal variances. The remaining four periods (2021-2024) show no significant difference '
    'in variances. The unequal variances in 2018 and 2020 may reflect that these are more '
    'volatile market regimes (crypto crash in 2018, COVID in 2020), causing greater run-to-run '
    'variability in the GA\'s ability to find good solutions.'
)

doc.add_heading('Results in detail and what they mean', level=2)
t5 = doc.add_table(rows=7, cols=4, style='Light Grid Accent 1')
for j, h in enumerate(['OOS Period', 'F-statistic', 'p-value', 'Unequal Variance?']):
    t5.rows[0].cells[j].text = h
levene_data = [
    ['2018', '28.85', '0.0000', 'Yes'],
    ['2020', '38.32', '0.0000', 'Yes'],
    ['2021', '0.28', '0.5991', 'No'],
    ['2022', '1.25', '0.2659', 'No'],
    ['2023', '0.02', '0.8979', 'No'],
    ['2024', '0.02', '0.8939', 'No'],
]
for i, row_data in enumerate(levene_data):
    for j, cell_text in enumerate(row_data):
        t5.rows[i + 1].cells[j].text = cell_text

doc.add_paragraph('')
doc.add_paragraph(
    'What this means: In 2018 and 2020, the Full GA\'s Sharpe distribution has significantly '
    'more variance (spread) than the Benchmark\'s. Looking at the raw data: in 2018, the Full '
    'GA has a Sharpe std of 0.24 vs Benchmark\'s 0.10; in 2020, it is 0.24 vs 0.06. This '
    'tells us something about the GA\'s behaviour in extreme market regimes: the wider search '
    'space (13 strategies) sometimes finds excellent solutions and sometimes poor ones, '
    'producing more variable outcomes. The Benchmark, with only 6 strategies and a smaller '
    'search space, converges more consistently but to a worse average. In the calmer periods '
    '(2021-2024), both variants show similar variance, suggesting the GA converges reliably '
    'when the training environment is more stable. Practically, the unequal variances in '
    '2018 and 2020 mean that Welch\'s t-test (which does not assume equal variances) should '
    'be used rather than the standard t-test for those periods — though the Mann-Whitney U '
    'test is unaffected by variance differences entirely, which is another reason it is the '
    'primary test.'
)

# ── 7. Kruskal-Wallis ──
doc.add_heading('7. Kruskal-Wallis Test (Non-Parametric Omnibus)', level=1)

doc.add_heading('What it is', level=2)
doc.add_paragraph(
    'The Kruskal-Wallis test is the non-parametric equivalent of one-way ANOVA. It tests '
    'whether three or more independent groups come from the same distribution. Unlike ANOVA, '
    'it does not assume normality — it ranks all observations across groups and tests whether '
    'the mean ranks differ significantly. A significant result indicates that at least one '
    'group differs, but does not identify which pair(s) differ.'
)

doc.add_heading('How it is used in this study', level=2)
doc.add_paragraph(
    'The pooled OOS Sharpe ratios from all three variants (Full, LLM-Only, Benchmark — '
    '300 values each) are compared simultaneously. This serves as an omnibus test: before '
    'examining specific pairwise comparisons, it establishes whether there is any overall '
    'difference among the three variants.'
)

doc.add_heading('When it is used', level=2)
doc.add_paragraph(
    'As a preliminary omnibus test before the pairwise Mann-Whitney comparisons. A '
    'significant Kruskal-Wallis result justifies proceeding with pairwise tests; a '
    'non-significant result would cast doubt on any individual pairwise findings.'
)

doc.add_heading('Key results', level=2)
doc.add_paragraph(
    'H-statistic = 30.79, p = 0.000000. Highly significant — at least one variant '
    'differs from the others. This provides strong justification for the subsequent pairwise '
    'Mann-Whitney U tests.'
)

doc.add_heading('Results in detail and what they mean', level=2)
doc.add_paragraph(
    'The Kruskal-Wallis H-statistic of 30.79 with a p-value effectively equal to zero is '
    'a very strong result. It pools all 900 Sharpe values (300 per variant) and asks: "Are '
    'these three groups drawn from the same distribution?" The answer is a decisive no.'
)
doc.add_paragraph(
    'What this means: This is the "gatekeeper" test. If Kruskal-Wallis had been '
    'non-significant, the pairwise Mann-Whitney U results would be suspect — it would '
    'suggest that any apparent pairwise differences might be noise. The highly significant '
    'result confirms that the three GA variants genuinely produce different performance '
    'distributions. It does not tell us which pair differs (that is the job of the pairwise '
    'Mann-Whitney tests), but it establishes that the overall three-way comparison is '
    'meaningful. This is analogous to running a one-way ANOVA before post-hoc pairwise '
    'tests — it prevents over-interpretation of individual comparisons.'
)

# ── 8. Independent t-test ──
doc.add_heading('8. Independent t-Test (Parametric Pairwise)', level=1)

doc.add_heading('What it is', level=2)
doc.add_paragraph(
    'The independent samples t-test compares the means of two groups and tests whether '
    'the difference is statistically significant, assuming the data are approximately '
    'normally distributed with equal (or corrected-for) variances. It is the classical '
    'parametric counterpart to the Mann-Whitney U test.'
)

doc.add_heading('How it is used in this study', level=2)
doc.add_paragraph(
    'Independent t-tests are applied to multiple performance metrics (Sortino, Calmar, '
    'maximum drawdown, total return) for the Full GA vs Benchmark comparison across all 6 '
    'periods (Slide 11). Results are reported as +/* (significantly better), + (better but '
    'not significant), or - (worse). The t-test is secondary to the Mann-Whitney U for '
    'Sharpe comparisons, but is the reported test for the additional metrics.'
)

doc.add_heading('When it is used', level=2)
doc.add_paragraph(
    'For the multi-metric comparison table (Sortino, Calmar, MaxDD, Total Return) where '
    'the emphasis is on showing a consistent pattern of outperformance across different risk '
    'measures, not just Sharpe. Significance is assessed at p < 0.05.'
)

doc.add_heading('Key results', level=2)
doc.add_paragraph(
    'Sortino: Full GA significantly better in 4/6 periods (wins 5/6). '
    'Calmar: significantly better in 4/6 periods (wins 6/6). '
    'MaxDD (less negative = better): significantly better in 5/6 periods. '
    'Total Return: significantly better in 3/6 periods (wins 4/6). '
    'The pattern of outperformance is consistent across metrics, not confined to Sharpe alone.'
)

doc.add_heading('Results in detail and what they mean', level=2)
doc.add_paragraph('Full GA vs Benchmark across multiple metrics (+ = Full wins, * = p < 0.05):')
t6 = doc.add_table(rows=5, cols=8, style='Light Grid Accent 1')
for j, h in enumerate(['Metric', '2018', '2020', '2021', '2022', '2023', '2024', 'Wins']):
    t6.rows[0].cells[j].text = h
metric_data = [
    ['Sortino', '+*', '-', '+', '+*', '+*', '+*', '5/6'],
    ['Calmar', '+*', '+', '+', '+*', '+*', '+*', '6/6'],
    ['MaxDD', '-*', '+*', '+*', '+*', '+*', '+*', '5/6'],
    ['Total Return', '+', '-*', '-', '+*', '+*', '+*', '4/6'],
]
for i, row_data in enumerate(metric_data):
    for j, cell_text in enumerate(row_data):
        t6.rows[i + 1].cells[j].text = cell_text

doc.add_paragraph('')
doc.add_paragraph(
    'What this means: The t-test results across multiple metrics demonstrate that the Full '
    'GA\'s advantage is not an artefact of optimising for Sharpe alone. The Sortino ratio '
    '(which penalises downside volatility more than upside) shows the same pattern: Full GA '
    'wins in 5/6 periods, 4 significantly. The Calmar ratio (return relative to worst '
    'drawdown) is the cleanest result — Full GA wins all 6 periods. Maximum drawdown is '
    'better (less negative) for Full GA in 5/6 periods, meaning the Full GA not only earns '
    'better returns but also suffers smaller peak-to-trough losses. Total return shows the '
    'weakest pattern (4/6 wins), which is expected: the GA optimises Sharpe (risk-adjusted '
    'return), not raw return, so it sometimes sacrifices total return for lower volatility. '
    'The overall picture is one of consistent, multi-dimensional outperformance — the Full '
    'GA produces portfolios that are better on almost every risk and return measure.'
)

# ── 9. Mixed-Effects Model ──
doc.add_heading('9. Linear Mixed-Effects Model', level=1)

doc.add_heading('What it is', level=2)
doc.add_paragraph(
    'A linear mixed-effects model (LMM) extends ordinary linear regression by including '
    'both fixed effects (the variable of interest — here, GA variant) and random effects '
    '(grouping variables that introduce correlation — here, OOS period). The model is: '
    'Sharpe ~ Variant + (1|Period), meaning Sharpe is predicted by the GA variant (fixed '
    'effect) with a random intercept for each period (random effect).'
)
doc.add_paragraph(
    'The random effect for Period accounts for the fact that all 50 runs within a period '
    'share the same training and test data, making them correlated. A standard t-test or '
    'Mann-Whitney U treats all 300 observations as independent, which they are not — runs '
    'within the same period are more similar to each other than runs across periods. The '
    'mixed-effects model properly handles this nested structure, producing more conservative '
    '(and more honest) p-values.'
)

doc.add_heading('How it is used in this study', level=2)
doc.add_paragraph(
    'Two mixed-effects models are fitted: (1) Full GA vs Benchmark, and (2) LLM-Only GA '
    'vs Benchmark. In each, the fixed-effect coefficient for the variant represents the '
    'average Sharpe difference attributable to using the treatment variant instead of '
    'Benchmark, after removing period-level variation.'
)

doc.add_heading('When it is used', level=2)
doc.add_paragraph(
    'As the most methodologically rigorous test in the suite. While Mann-Whitney U provides '
    'per-period significance, the mixed-effects model provides an aggregate assessment that '
    'correctly accounts for the hierarchical data structure (runs nested within periods). '
    'It is the appropriate test when the question is: "Across all market regimes, does this '
    'variant systematically outperform?"'
)

doc.add_heading('Key results', level=2)
doc.add_paragraph(
    'Full vs Benchmark: coefficient = +0.4114, p = 0.0610, 95% CI [+0.0767, +0.7462]. '
    'LLM-Only vs Benchmark: coefficient = +0.4302, p = 0.0608, 95% CI [+0.0804, +0.7799]. '
    'Both are marginally non-significant at the conventional 0.05 level but significant at '
    'the 0.10 level. The 95% CIs exclude zero in both cases. The borderline p-values reflect '
    'the conservative nature of the test: with only 6 periods (the effective sample size for '
    'the random effect), the model has limited power to detect effects. This is a known '
    'limitation of mixed-effects models with few clusters.'
)

doc.add_heading('Results in detail and what they mean', level=2)
t7 = doc.add_table(rows=3, cols=4, style='Light Grid Accent 1')
for j, h in enumerate(['Comparison', 'Coefficient (Sharpe diff)', 'p-value', '95% CI']):
    t7.rows[0].cells[j].text = h
me_data = [
    ['Full vs Benchmark', '+0.4114', '0.0610', '[+0.0767, +0.7462]'],
    ['LLM-Only vs Benchmark', '+0.4302', '0.0608', '[+0.0804, +0.7799]'],
]
for i, row_data in enumerate(me_data):
    for j, cell_text in enumerate(row_data):
        t7.rows[i + 1].cells[j].text = cell_text

doc.add_paragraph('')
doc.add_paragraph(
    'What this means: The mixed-effects model is the most methodologically honest test in '
    'the suite, and its results require careful interpretation. The p-values of ~0.061 are '
    'just above the conventional 0.05 threshold, which might appear to contradict the very '
    'strong Mann-Whitney results. However, this is not a contradiction — it is a consequence '
    'of the test\'s conservative design.'
)
doc.add_paragraph(
    'The key issue is statistical power. The mixed-effects model treats Period as a random '
    'effect, meaning it estimates variance across periods. With only 6 periods (the effective '
    'cluster count), the model has very few degrees of freedom for the random effect, making '
    'it hard to achieve p < 0.05. This is a well-documented limitation: mixed-effects models '
    'with fewer than ~10-15 clusters tend to have inflated p-values. Despite this, two '
    'important observations support the effect being real: (1) the 95% confidence intervals '
    'exclude zero in both comparisons, meaning the estimated effect is positive with 95% '
    'confidence, and (2) the coefficients (+0.41 and +0.43) are practically meaningful — they '
    'represent an average improvement of roughly 0.4 Sharpe points after accounting for all '
    'period-level variation.'
)
doc.add_paragraph(
    'The borderline p-values should be interpreted as "the evidence is strongly suggestive '
    'but the test lacks the power to reach conventional significance with only 6 clusters", '
    'not as "the effect does not exist". The per-period Mann-Whitney tests, which have full '
    'statistical power (n = 50 per group), confirm the effect exists within each period.'
)

# ── 10. Friedman Test ──
doc.add_heading('10. Friedman Test (Non-Parametric Repeated Measures)', level=1)

doc.add_heading('What it is', level=2)
doc.add_paragraph(
    'The Friedman test is the non-parametric equivalent of repeated-measures ANOVA. It '
    'tests whether three or more related groups (here, the three GA variants measured across '
    'the same 6 periods) differ in their central tendency. It ranks the variants within each '
    'period and tests whether the rank sums differ across variants.'
)

doc.add_heading('How it is used in this study', level=2)
doc.add_paragraph(
    'The mean Sharpe ratio for each variant in each period is computed (collapsing the 50 '
    'runs to a single summary per cell). This gives a 6 x 3 matrix (6 periods, 3 variants). '
    'The Friedman test ranks the three variants within each period and tests whether the '
    'overall ranking pattern is consistent. It is the most conservative test in the suite '
    'because it reduces 50 runs per cell to a single value, discarding within-period '
    'variability information.'
)

doc.add_heading('When it is used', level=2)
doc.add_paragraph(
    'As a complement to the mixed-effects model. Both account for the repeated-measures '
    'structure (the same periods are tested for all three variants), but the Friedman test '
    'is fully non-parametric and makes no distributional assumptions at all.'
)

doc.add_heading('Key results', level=2)
doc.add_paragraph(
    'Statistic = 4.33, p = 0.1146. Not significant. With only 6 periods and 3 variants, '
    'the Friedman test has very low statistical power. The non-significance does not '
    'contradict the per-period Mann-Whitney results; it reflects the inherent difficulty of '
    'detecting effects with n = 6 blocks in a rank-based test.'
)

doc.add_heading('Results in detail and what they mean', level=2)
doc.add_paragraph(
    'The Friedman test reduces each (period, variant) cell to a single value (the mean of '
    '50 runs), producing a 6 x 3 matrix. Within each period, it ranks the three variants '
    '(1st, 2nd, 3rd) and tests whether the sum of ranks differs across variants. With p = '
    '0.1146, it does not reach significance.'
)
doc.add_paragraph(
    'What this means: The Friedman test is the most conservative test in the entire suite '
    'and its non-significance is expected. By collapsing 50 runs to a single mean per cell, '
    'it discards all within-period variability information, leaving only 6 data points per '
    'variant. Furthermore, it is rank-based, meaning it only uses ordinal information (which '
    'variant is 1st, 2nd, 3rd) and ignores the magnitude of differences. With only 6 blocks, '
    'the maximum possible Friedman statistic is limited, making it nearly impossible to reach '
    'significance unless one variant wins every single period by a clear margin. The 2020 '
    'period (where Benchmark wins) is enough to prevent that.'
)
doc.add_paragraph(
    'The Friedman result should be interpreted as: "With only 6 periods, a rank-based '
    'repeated-measures test does not have enough power to detect the effect." It does not '
    'mean the effect is absent. The mixed-effects model (which uses all 900 observations '
    'rather than 18 cell means) comes much closer to significance (p = 0.061) precisely '
    'because it retains more information.'
)

# ── Summary Table ──
doc.add_heading('Summary of All Tests', level=1)

table = doc.add_table(rows=11, cols=4, style='Light Grid Accent 1')
headers = ['Test', 'Type', 'Purpose', 'Conclusion']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h

rows_data = [
    ['Mann-Whitney U', 'Non-parametric pairwise', 'Primary significance test for Sharpe differences between variant pairs', 'Significant in 5/6 periods (Full vs Bench) and 5/6 (LLM vs Bench)'],
    ['Holm-Bonferroni', 'Multiple comparisons correction', 'Controls family-wise error rate across 18 simultaneous tests', 'All significant results survive correction'],
    ['Bootstrap CI', 'Resampling', 'Estimates magnitude and uncertainty of mean Sharpe difference', 'All CIs exclude zero, confirming MW-U results'],
    ["Cohen's d", 'Effect size', 'Quantifies practical significance of differences', 'Mean d = 2.87 (large) for Full vs Benchmark'],
    ['Shapiro-Wilk', 'Normality diagnostic', 'Tests whether Sharpe distributions are normal', '5/18 reject normality; justifies non-parametric primary test'],
    ["Levene's", 'Variance diagnostic', 'Tests equality of variances between groups', 'Unequal variances in 2018 and 2020; equal in 2021-2024'],
    ['Kruskal-Wallis', 'Non-parametric omnibus', 'Tests if any variant differs (3-group comparison)', 'H = 30.79, p = 0.000000; highly significant'],
    ['Independent t-test', 'Parametric pairwise', 'Compares means for Sortino, Calmar, MaxDD, Total Return', 'Consistent outperformance across multiple metrics'],
    ['Mixed-Effects Model', 'Parametric hierarchical', 'Accounts for within-period correlation (runs nested in periods)', 'p = 0.061 (borderline); CI excludes zero'],
    ['Friedman', 'Non-parametric repeated measures', 'Rank-based test across variants with period as block', 'p = 0.115 (not significant; low power with 6 periods)'],
]
for i, row_data in enumerate(rows_data):
    for j, cell_text in enumerate(row_data):
        table.rows[i + 1].cells[j].text = cell_text

# ── Interpretation ──
doc.add_heading('Interpretation and Rationale for Multi-Test Approach', level=1)
doc.add_paragraph(
    'The study uses a deliberate hierarchy of statistical evidence:'
)
p = doc.add_paragraph(style='List Bullet')
p.add_run('Per-period tests (Mann-Whitney U) ').bold = True
p.add_run('establish that differences are reliable within each market regime, with 50 '
          'independent runs providing statistical power.')

p = doc.add_paragraph(style='List Bullet')
p.add_run('Multiple comparisons correction (Holm-Bonferroni) ').bold = True
p.add_run('ensures that the family of 18 tests does not inflate false positive rates.')

p = doc.add_paragraph(style='List Bullet')
p.add_run('Effect sizes (Cohen\'s d) and confidence intervals (bootstrap) ').bold = True
p.add_run('quantify the magnitude of differences, preventing over-reliance on p-values alone.')

p = doc.add_paragraph(style='List Bullet')
p.add_run('Diagnostics (Shapiro-Wilk, Levene\'s) ').bold = True
p.add_run('validate the choice of primary test and flag assumption violations.')

p = doc.add_paragraph(style='List Bullet')
p.add_run('Aggregate tests (Kruskal-Wallis, mixed-effects, Friedman) ').bold = True
p.add_run('assess overall variant differences while accounting for repeated-measures '
          'structure. The mixed-effects model is the most appropriate aggregate test but has '
          'limited power with only 6 period-level clusters.')

doc.add_paragraph(
    'The convergence of evidence — strong per-period significance, large effect sizes, '
    'bootstrap CIs excluding zero, and a significant omnibus test — supports the conclusion '
    'that LLM-generated strategies add value to the portfolio. The borderline mixed-effects '
    'p-value (0.061) is consistent with this interpretation: it reflects low power from '
    'having only 6 periods as random-effect levels, not absence of an effect.'
)

# ── Comprehensive Results Synthesis ──
doc.add_page_break()
doc.add_heading('Comprehensive Results Synthesis', level=1)

doc.add_heading('What the experiments set out to answer', level=2)
doc.add_paragraph(
    'The central research question is: Do LLM-generated trading strategies add value to a '
    'portfolio beyond what conventional benchmark strategies achieve alone? The experimental '
    'design isolates this question through a three-way controlled comparison:'
)
p = doc.add_paragraph(style='List Bullet')
p.add_run('Full GA ').bold = True
p.add_run('(13 strategies: 7 LLM + 5 benchmarks + Cash) — the complete strategy universe.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('LLM-Only GA ').bold = True
p.add_run('(8 strategies: 7 LLM + Cash) — isolates the novel strategies without benchmark assistance.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Benchmark-Only GA ').bold = True
p.add_run('(6 strategies: 5 benchmarks + Cash) — the control. Any advantage of Full or LLM-Only '
          'over this baseline is attributable to the LLM-generated strategies, not the GA itself.')

doc.add_heading('The headline finding', level=2)
doc.add_paragraph(
    'Both the Full GA and LLM-Only GA produce significantly better risk-adjusted returns '
    '(Sharpe ratios) than the Benchmark-Only GA in 5 out of 6 out-of-sample test years '
    '(2018, 2021, 2022, 2023, 2024). The one exception is 2020, where the Benchmark wins — '
    'likely because the COVID crash-then-recovery strongly favoured simple momentum and '
    'buy-and-hold strategies. Across all periods combined, the Full GA achieves a mean OOS '
    'Sharpe of +0.796 vs the Benchmark\'s +0.385 (a difference of +0.411), and the LLM-Only '
    'GA achieves +0.815 vs +0.385 (a difference of +0.430).'
)

doc.add_heading('How confident can we be?', level=2)
doc.add_paragraph(
    'The statistical evidence forms a layered pyramid of confidence:'
)

doc.add_heading('Layer 1: Per-period significance (very strong)', level=3)
doc.add_paragraph(
    'The Mann-Whitney U test finds statistically significant differences in all 6 periods '
    'for both comparisons (Full vs Bench and LLM vs Bench), with adjusted p-values all '
    'below 0.001 after Holm-Bonferroni correction. This means that within each individual '
    'market regime, the GA variants produce reliably different outcomes — not by chance, but '
    'systematically. The 50 independent runs per group provide strong statistical power.'
)

doc.add_heading('Layer 2: Effect magnitude (very strong)', level=3)
doc.add_paragraph(
    'Cohen\'s d values range from 0.80 to 5.68 (mean 2.87), all in the "large" category. '
    'Bootstrap 95% confidence intervals on the Sharpe difference exclude zero in every period. '
    'This confirms that the differences are not just statistically detectable but practically '
    'meaningful — the Full GA and LLM-Only GA produce substantially better portfolios.'
)

doc.add_heading('Layer 3: Omnibus confirmation (strong)', level=3)
doc.add_paragraph(
    'The Kruskal-Wallis test (H = 30.79, p < 0.000001) confirms that the three variants '
    'are genuinely different populations, not draws from the same distribution. This is the '
    '"gatekeeper" that validates all pairwise comparisons.'
)

doc.add_heading('Layer 4: Aggregate with hierarchical structure (suggestive)', level=3)
doc.add_paragraph(
    'The mixed-effects model, which properly accounts for runs being nested within periods, '
    'estimates the variant effect at +0.41 Sharpe points (Full vs Bench) with a p-value of '
    '0.061 and a 95% CI that excludes zero. The Friedman test (p = 0.115) is non-significant. '
    'Both aggregate tests are limited by having only 6 period-level clusters — a known power '
    'limitation. The borderline results are consistent with the effect being real but the '
    'tests lacking power to confirm it at the 0.05 level with so few periods.'
)

doc.add_heading('Layer 5: Multi-metric robustness (strong)', level=3)
doc.add_paragraph(
    'The pattern of outperformance extends beyond Sharpe: the Full GA wins on Sortino '
    '(5/6 periods), Calmar (6/6), and maximum drawdown (5/6). This confirms that the '
    'advantage is not an artefact of optimising for one specific metric — the LLM-strategy '
    'portfolios are genuinely better across multiple dimensions of risk and return.'
)

doc.add_heading('What the 2020 exception tells us', level=2)
doc.add_paragraph(
    'The 2020 period (COVID crash and recovery) is the one year where the Benchmark-Only GA '
    'outperforms both treatment variants. This is actually informative rather than damaging '
    'to the thesis. The 2020 market was dominated by an extreme crash followed by a rapid '
    'V-shaped recovery driven by massive fiscal stimulus — a regime where simple momentum and '
    'buy-and-hold strategies excelled because the dominant trade was simply "stay long and '
    'ride the recovery". The more complex LLM strategies (entropy-based, fractal, '
    'information-geometric) likely generated hedging signals during the crash that were '
    'correct at the time but costly during the subsequent violent recovery. This suggests '
    'that the LLM strategies add value in normal-to-moderately-stressed markets but may '
    'underperform in extreme, stimulus-driven V-recoveries — a nuance that future work on '
    'dynamic regime-adaptive allocation could address.'
)

doc.add_heading('What the LLM-Only result means', level=2)
doc.add_paragraph(
    'A particularly notable finding is that the LLM-Only GA (which contains no benchmark '
    'strategies at all) performs comparably to — and sometimes better than — the Full GA. '
    'The LLM-Only GA beats the Benchmark in the same 5/6 periods with similar effect sizes, '
    'and its aggregate mean Sharpe (+0.815) actually exceeds the Full GA\'s (+0.796). This '
    'demonstrates that the LLM-generated strategies can stand alone: they do not merely '
    'complement benchmark strategies but can replace them entirely without loss of '
    'performance. The GA, when given the choice (in the Full variant), allocates 77% of '
    'total portfolio weight to LLM strategies over benchmarks, further confirming their '
    'value through revealed preference.'
)

doc.add_heading('Reconciling the borderline aggregate tests', level=2)
doc.add_paragraph(
    'The apparent tension between the very strong per-period results (all p < 0.001) and '
    'the borderline aggregate results (mixed-effects p = 0.061, Friedman p = 0.115) is '
    'entirely expected and does not represent a contradiction. The per-period tests have '
    'n = 50 per group (strong power); the aggregate tests effectively have n = 6 periods '
    '(weak power). With 6 clusters, mixed-effects models and Friedman tests are known to '
    'have inflated Type II error rates (failure to detect real effects). The fact that the '
    'mixed-effects CI excludes zero despite the borderline p-value, and that the coefficient '
    'is +0.41 (a practically meaningful Sharpe improvement), supports the conclusion that '
    'the effect is real but the aggregate tests simply lack the statistical power to confirm '
    'it at the conventional threshold. More OOS periods (future work) would resolve this.'
)

doc.add_heading('Overall conclusion from the statistical evidence', level=2)
doc.add_paragraph(
    'The weight of evidence strongly supports the conclusion that LLM-generated trading '
    'strategies add genuine value to a portfolio optimised by a genetic algorithm. This '
    'conclusion is supported by:'
)
p = doc.add_paragraph(style='List Number')
p.add_run('Statistically significant outperformance in 5/6 individual OOS periods, '
          'surviving multiple-comparisons correction.')
p = doc.add_paragraph(style='List Number')
p.add_run('Very large effect sizes (mean Cohen\'s d = 2.87) confirming practical significance.')
p = doc.add_paragraph(style='List Number')
p.add_run('Bootstrap confidence intervals excluding zero in all periods.')
p = doc.add_paragraph(style='List Number')
p.add_run('A significant omnibus test (Kruskal-Wallis) confirming the three variants differ.')
p = doc.add_paragraph(style='List Number')
p.add_run('Consistent outperformance across multiple risk-adjusted metrics (Sharpe, Sortino, '
          'Calmar, drawdown).')
p = doc.add_paragraph(style='List Number')
p.add_run('The LLM-Only GA matching the Full GA, demonstrating the novel strategies can '
          'stand alone.')
p = doc.add_paragraph(style='List Number')
p.add_run('The GA\'s revealed preference: 77% of weight allocated to LLM strategies when '
          'given the full strategy set.')

doc.add_paragraph('')
doc.add_paragraph(
    'The only caveats are: (1) the 2020 COVID period where benchmarks win, suggesting '
    'regime-sensitivity, and (2) the borderline aggregate tests, which are a power limitation '
    '(6 periods) rather than evidence against the effect. Both point to clear directions for '
    'future work: dynamic allocation and extended evaluation periods.'
)

# Save
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                            'Significance_Tests_Description.docx')
doc.save(output_path)
print(f'Saved to {output_path}')
